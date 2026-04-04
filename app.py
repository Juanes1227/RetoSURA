
import openpyxl
import streamlit as st
import pandas as pd
import numpy as np
import os
import io
import google.generativeai as genai
from docx import Document


# 1. Configuración básica de Streamlit
st.set_page_config(page_title="Cotizador de Renovaciones", page_icon="🛡️", layout="wide")
st.title("🛡️ Herramienta de Tarifación para Renovaciones")

# --- SISTEMA DE AUTO-GUARDADO Y RECUPERACIÓN ---
archivos_respaldo = {
    'df_polizas': 'backup_polizas.pkl',
    'df_curva': 'backup_curva.pkl',
    'df_siniestros': 'backup_siniestros.pkl',
    'datos_procesados': 'backup_procesados.pkl'
}

for key, archivo in archivos_respaldo.items():
    if key not in st.session_state:
        if os.path.exists(archivo):
            st.session_state[key] = pd.read_pickle(archivo)
        else:
            st.session_state[key] = None

if 'carta_generada' not in st.session_state:
    st.session_state['carta_generada'] = None
if 'poliza_actual_ia' not in st.session_state:
    st.session_state['poliza_actual_ia'] = None

# 3. Función de procesamiento principal (MODIFICADA CON TPR POR PRODUCTO)
@st.cache_data(show_spinner=False)
def procesar_datos(polizas_raw, curva_raw, siniestros_raw):
    polizas = polizas_raw.copy()
    curva_riesgo = curva_raw.copy()
    siniestros = siniestros_raw.copy()
    
    curva_siniestros = pd.merge(siniestros, polizas, on='Poliza_Id', how='left')
    
    # --- NUEVO: TPR POR PRODUCTO (2025) ---
    curva_2025 = curva_siniestros[curva_siniestros["Year"] == 2025]
    tpr_por_producto = curva_2025.groupby('Producto_Desc').agg(
        Total_Siniestros=('Valor del Siniestro', 'sum'),
        Total_VA=('VA_Ajuste_Prueba', 'sum')
    ).reset_index()
    
    tpr_por_producto['Total_Siniestros'] = tpr_por_producto['Total_Siniestros'] / 100 
    tpr_por_producto['TPR'] = tpr_por_producto['Total_Siniestros'] / tpr_por_producto['Total_VA']
    # ---------------------------------------

    polizas_riesgo = pd.merge(polizas, curva_riesgo, left_on=['Amparo_Id', 'Edad_Cliente'], right_on=['Amparo_ID', 'Edad'], how='inner')
    polizas_riesgo["PPRC"] = polizas_riesgo["VA_Ajuste_Prueba"] * polizas_riesgo["Tasa"]
    polizas_riesgo["TCPC"] = polizas_riesgo["Tasa"] / (1 - 0.4)
    polizas_riesgo['VA_Solo_Vida'] = np.where(polizas_riesgo['Amparo_Desc'].str.startswith('VIDA', na=False), polizas_riesgo['VA_Ajuste_Prueba'], 0)

    tasa_unica_poliza = polizas_riesgo.groupby(['Poliza_Id'])[['PPRC', "VA_Solo_Vida"]].sum().reset_index()
    tasa_unica_poliza = tasa_unica_poliza.rename(columns={'PPRC': 'Prima por Cobertura'})
    tasa_unica_poliza["Tasa Unica pura de riesgo"] = tasa_unica_poliza["Prima por Cobertura"] / tasa_unica_poliza["VA_Solo_Vida"]
    tasa_unica_poliza["Tasa Unica comercial"] = tasa_unica_poliza["Tasa Unica pura de riesgo"] / (1 - 0.4)

    siniestros["n"] = siniestros.groupby(['Poliza_Id'])["Poliza_Id"].transform('count')
    siniestros['Valor_Calculo'] = np.where(siniestros['Valor del Siniestro'] > 2350, 2350, siniestros['Valor del Siniestro'])
    siniestros["varianza_por_cliente"] = siniestros.groupby('Poliza_Id')['Valor_Calculo'].transform('var', ddof=1).fillna(0)
    siniestros["media_por_cliente"] = siniestros.groupby('Poliza_Id')['Valor_Calculo'].transform('mean').fillna(0)

    siniestros_gb = siniestros.groupby(['Poliza_Id']).agg({
        'Valor del Siniestro': 'mean', 'n': 'mean', 'varianza_por_cliente': 'mean', 'media_por_cliente': 'mean'
    }).reset_index()

    curva_siniestros["Valor_Calculo"] = curva_siniestros["Valor del Siniestro"].clip(upper=181418.13)
    siniestros_gb_ = curva_siniestros.groupby('Poliza_Id')['Valor_Calculo'].agg(['mean', 'var']).fillna(0)

    EPV = siniestros_gb_['var'].mean()
    VHM = siniestros_gb_['mean'].var(ddof=1)
    K_nuevo = EPV / VHM if VHM != 0 else 0

    siniestros_gb["K"] = K_nuevo
    siniestros_gb["Z"] = (siniestros_gb["n"] / (siniestros_gb["n"] + siniestros_gb["K"]))

    # --- NUEVO: CRUCE DE TPR CON CREDIBILIDAD ---
    # 1. Recuperamos el Producto_Desc para siniestros_gb desde pólizas
    siniestros_gb = siniestros_gb.merge(polizas[['Poliza_Id', 'Producto_Desc']].drop_duplicates(), on='Poliza_Id', how='left')
    
    # 2. Traemos la TPR por producto
    siniestros_gb = siniestros_gb.merge(tpr_por_producto[['Producto_Desc', 'TPR']], on='Producto_Desc', how='left')

    # 3. Consolidamos en tasa_unica_poliza (Usando how='left' para no perder pólizas sin reclamos)
    tasa_unica_poliza = pd.merge(tasa_unica_poliza, siniestros_gb[['Poliza_Id', 'Z', 'TPR']], on='Poliza_Id', how='left')
    
    # Limpiamos vacíos para pólizas sin siniestros/productos nuevos
    tasa_unica_poliza["Z"] = tasa_unica_poliza["Z"].fillna(0)
    
    # Rellenamos TPR nulas con la TPR global promedio del 2025 como contingencia
    tpr_global = tpr_por_producto['TPR'].mean() if not tpr_por_producto.empty else 0
    tasa_unica_poliza["TPR"] = tasa_unica_poliza["TPR"].fillna(tpr_global)

    # Cálculo Final
    tasa_unica_poliza["tasa_cred"] = (tasa_unica_poliza["Z"] * tasa_unica_poliza["Tasa Unica pura de riesgo"]) + ((1 - tasa_unica_poliza["Z"]) * tasa_unica_poliza["TPR"])
    tasa_unica_poliza["tasa_cred_com"] = tasa_unica_poliza["tasa_cred"] / (1 - 0.4)
    tasa_unica_poliza["prima_recomendada"] = tasa_unica_poliza["tasa_cred"] * tasa_unica_poliza["VA_Solo_Vida"]

    return tasa_unica_poliza


# 4. BARRA LATERAL: Configuración y Datos
st.sidebar.header("🔑 Configuración IA")
try:
    API_KEY = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=API_KEY)
    modelo = genai.GenerativeModel('gemini-2.5-flash')
except KeyError:
    st.sidebar.warning("API Key no encontrada en secrets.toml")
    modelo = None

st.sidebar.divider()
st.sidebar.header("📂 1. Gestión de Bases de Datos")

estado_p = "🟢" if st.session_state['df_polizas'] is not None else "🔴"
estado_c = "🟢" if st.session_state['df_curva'] is not None else "🔴"
estado_s = "🟢" if st.session_state['df_siniestros'] is not None else "🔴"
st.sidebar.markdown(f"{estado_p} Pólizas | {estado_c} Curva | {estado_s} Siniestros")

with st.sidebar.expander("📄 Cargar / Actualizar Pólizas"):
    file_polizas = st.file_uploader("Archivo de Pólizas (.txt, .csv)", type=['txt', 'csv'], key="up_polizas")
    modo_carga = st.radio("Acción a realizar:", ["Reemplazar toda la base", "Actualizar / Agregar específicas"])
    
    if st.button("Aplicar Pólizas") and file_polizas:
        df_nuevo = pd.read_csv(file_polizas, sep='|', decimal=",")
        if modo_carga == "Reemplazar toda la base" or st.session_state['df_polizas'] is None:
            st.session_state['df_polizas'] = df_nuevo
        else:
            ids_nuevos = df_nuevo['Poliza_Id'].unique()
            df_actual = st.session_state['df_polizas']
            df_actual = df_actual[~df_actual['Poliza_Id'].isin(ids_nuevos)]
            st.session_state['df_polizas'] = pd.concat([df_actual, df_nuevo], ignore_index=True)
            
        st.session_state['df_polizas'].to_pickle(archivos_respaldo['df_polizas'])
        st.success("Base de pólizas guardada.")

with st.sidebar.expander("📈 Cargar / Actualizar Curva de Riesgo"):
    file_curva = st.file_uploader("Archivo Curva (.xlsx)", type=['xlsx'], key="up_curva")
    if st.button("Aplicar Curva") and file_curva:
        st.session_state['df_curva'] = pd.read_excel(file_curva)
        st.session_state['df_curva'].to_pickle(archivos_respaldo['df_curva'])
        st.success("Curva de riesgo guardada.")

with st.sidebar.expander("💥 Cargar / Actualizar Siniestros"):
    file_siniestros = st.file_uploader("Archivo Siniestros (.xlsx)", type=['xlsx'], key="up_siniestros")
    if st.button("Aplicar Siniestros") and file_siniestros:
        st.session_state['df_siniestros'] = pd.read_excel(file_siniestros)
        st.session_state['df_siniestros'].to_pickle(archivos_respaldo['df_siniestros'])
        st.success("Siniestros históricos guardados.")

st.sidebar.divider()

if st.session_state['df_polizas'] is not None and st.session_state['df_curva'] is not None and st.session_state['df_siniestros'] is not None:
    if st.sidebar.button("⚙️ Procesar y Calcular Tarifas", type="primary"):
        with st.spinner("Ejecutando modelos de credibilidad..."):
            try:
                df_final = procesar_datos(
                    st.session_state['df_polizas'], 
                    st.session_state['df_curva'], 
                    st.session_state['df_siniestros']
                )
                st.session_state['datos_procesados'] = df_final
                df_final.to_pickle(archivos_respaldo['datos_procesados'])
                st.sidebar.success("✅ Cálculos procesados y guardados.")
            except Exception as e:
                st.sidebar.error(f"Error al calcular: {e}")

if st.sidebar.button("🗑️ Borrar toda la base de datos"):
    for key, archivo in archivos_respaldo.items():
        st.session_state[key] = None
        if os.path.exists(archivo):
            os.remove(archivo)
    st.rerun()

# 5. PANTALLA PRINCIPAL: Buscador y Resultados
if st.session_state['datos_procesados'] is not None:
    df_resultados = st.session_state['datos_procesados']
    
    st.subheader("🔍 2. Consultar Póliza")
    
    col_busqueda1, col_busqueda2 = st.columns(2)
    with col_busqueda1:
        busqueda_texto = st.text_input("Buscar por ID de Póliza (Escriba y presione Enter):", placeholder="Ej: 109553448")
    with col_busqueda2:
        lista_polizas = df_resultados['Poliza_Id'].astype(str).unique()
        busqueda_lista = st.selectbox("O seleccione de la lista:", [""] + list(lista_polizas))

    poliza_seleccionada = busqueda_texto if busqueda_texto != "" else busqueda_lista

    if poliza_seleccionada:
        if poliza_seleccionada in lista_polizas:
            
            if st.session_state['poliza_actual_ia'] != poliza_seleccionada:
                st.session_state['carta_generada'] = None
                st.session_state['poliza_actual_ia'] = poliza_seleccionada

            datos_poliza = df_resultados[df_resultados['Poliza_Id'].astype(str) == poliza_seleccionada].iloc[0]
            prima_anterior = datos_poliza['Prima por Cobertura']
            prima_nueva = datos_poliza['prima_recomendada']
            
            st.header(f"Resultados para la Póliza: `{poliza_seleccionada}`")

            # --- Tarjetas Visuales Actualizadas ---
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric(label="TPR del Producto (2025)", value=f"{datos_poliza['TPR']:.5f}")
            with col2:
                st.metric(label="Factor Credibilidad (Z)", value=f"{datos_poliza['Z'] * 100:.2f}%")
            with col3:
                diferencia_tasa = datos_poliza['tasa_cred'] - datos_poliza['TPR']
                st.metric(label="Tasa Credibilizada", value=f"{datos_poliza['tasa_cred']:.5f}", 
                          delta=f"{diferencia_tasa:.5f} vs TPR", delta_color="inverse")
            with col4:
                st.metric(label="Tasa Cred. Comercial", value=f"{datos_poliza['tasa_cred_com']:.5f}")

            st.divider()

            st.subheader("Comparativo Financiero")
            col_a, col_b, col_c = st.columns(3)
            with col_a:
                st.info(f"**Valor Asegurado (Solo Vida):**\n\n${datos_poliza['VA_Solo_Vida']:,.2f}")
            with col_b:
                st.warning(f"**Prima Actual (Por Cobertura):**\n\n${datos_poliza['Prima por Cobertura']:,.2f}")
            with col_c:
                st.success(f"**PRIMA RECOMENDADA NUEVA:**\n\n${datos_poliza['prima_recomendada']:,.2f}")

            with st.expander("Ver tabla completa de la póliza (Detalles crudos)"):
                st.dataframe(pd.DataFrame(datos_poliza).T, use_container_width=True)
            
            # --- SECCIÓN DE IA ---
            st.divider()
            st.subheader("🤖 Asistente de Renovación IA")

            variacion_pct = ((prima_nueva - prima_anterior) / prima_anterior) * 100 if prima_anterior > 0 else 0
            st.info(f"Variación de la prima: **{variacion_pct:,.2f}%**")

            umbral_atipico = st.slider("Generar alerta si el aumento supera el (%):", 0, 100, 20)
            es_atipica = variacion_pct > umbral_atipico

            if es_atipica:
                st.warning("⚠️ Esta póliza presenta un incremento atípico. Se sugiere generar la carta de renovación.")

            if st.button("✨ Generar Carta al Cliente con IA"):
                if modelo is None:
                    st.error("Falta la configuración de la IA en los secretos.")
                else:
                    with st.spinner("Redactando propuesta comercial..."):
                        prompt = f"""
                        Actúa como un suscriptor de seguros experto y empático.
                        Redacta UNA carta formal dirigida al cliente de la póliza {poliza_seleccionada}.
                        
                        Contexto:
                        - Valor Asegurado: ${datos_poliza['VA_Solo_Vida']:,.2f}
                        - Prima Anterior: ${prima_anterior:,.2f}
                        - Prima Nueva Recomendada: ${prima_nueva:,.2f}
                        
                        Instrucciones:
                        Justifica suavemente el incremento de la prima mencionando el ajuste por riesgo y las condiciones macroeconómicas, sin usar jerga actuarial compleja. 
                        El tono debe ser profesional, agradecido por su fidelidad y enfocado en la protección continua que le brinda la póliza.
                        No incluyas saludos genéricos como [Nombre del Cliente], usa "Estimado Cliente".
                        """
                        respuesta = modelo.generate_content(prompt)
                        st.session_state['carta_generada'] = respuesta.text

            if st.session_state['carta_generada']:
                st.markdown("### Vista Previa de la Carta:")
                st.write(st.session_state['carta_generada'])
                
                doc = Document()
                doc.add_heading(f'Propuesta de Renovación - Póliza {poliza_seleccionada}', 0)
                doc.add_paragraph(st.session_state['carta_generada'])
                
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                st.download_button(
                    label="📥 Descargar Carta en Word (.docx)",
                    data=buffer,
                    file_name=f"Renovacion_Poliza_{poliza_seleccionada}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )

        else:
            st.error(f"❌ La póliza '{poliza_seleccionada}' no se encuentra en la base de datos actual.")
else:
    st.info("👈 Sube los archivos en la barra lateral y presiona 'Procesar y Calcular Tarifas' para comenzar.")